import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import os

# Configuración optimizada para móviles
st.set_page_config(page_title="Certificador Pro", layout="centered")

# --- CONFIGURACIÓN DE ARCHIVOS FIJOS ---
PLANTILLAS_CONFIG = [
    "certificacion.docx",
    "anexo.docx"
]

def extraer_etiquetas_fijas(nombres_archivos):
    """Lee los archivos locales y extrae las etiquetas {{etiqueta}}"""
    etiquetas = set()
    for nombre in nombres_archivos:
        if os.path.exists(nombre):
            try:
                doc = Document(nombre)
                # Buscar en párrafos
                for p in doc.paragraphs:
                    encontrados = re.findall(r"\{\{(.*?)\}\}", p.text)
                    for e in encontrados: etiquetas.add(e.strip())
                # Buscar en tablas
                for t in doc.tables:
                    for f in t.rows:
                        for c in f.cells:
                            encontrados = re.findall(r"\{\{(.*?)\}\}", c.text)
                            for e in encontrados: etiquetas.add(e.strip())
            except Exception as e:
                st.error(f"Error leyendo {nombre}: {e}")
    return sorted(list(etiquetas))

def aplicar_estilo_y_reemplazo(paragraph, datos, fuente_size):
    """
    Reemplaza etiquetas en un párrafo manteniendo Arial y el tamaño indicado,
    aplicando negrita a campos específicos.
    """
    texto_original = paragraph.text
    if "{{" not in texto_original:
        # Si no hay etiquetas, solo aseguramos la fuente Arial y el tamaño
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(fuente_size)
        return

    # Lista de etiquetas que deben ir en negrita
    tags_negrita = [
        "dirigido_a_la_institucion", 
        "nombre", 
        "cedula", 
        "actividad_economica_del_cliente"
    ]
    
    # Expresión regular para encontrar las etiquetas
    parts = re.split(r'(\{\{.*?\}\})', texto_original)
    
    # Limpiamos el párrafo para reconstruirlo con el nuevo formato
    paragraph.text = ""
    
    for part in parts:
        run = paragraph.add_run()
        # Verificar si la parte es una etiqueta {{etiqueta}}
        tag_match = re.match(r'\{\{(.*?)\}\}', part)
        
        if tag_match:
            tag_key = tag_match.group(1).strip()
            valor = datos.get(tag_key, part)
            run.text = valor
            # Aplicar negrita si está en la lista
            if tag_key in tags_negrita:
                run.bold = True
        else:
            run.text = part
            
        # Aplicar Arial y el tamaño correspondiente
        run.font.name = 'Arial'
        run.font.size = Pt(fuente_size)

def generar_documento(nombre_plantilla, datos):
    """Aplica los cambios a una plantilla con Arial y negritas selectivas"""
    doc = Document(nombre_plantilla)
    
    # Definir tamaño de fuente: 12 para anexo, 8 para el resto
    fuente_size = 12 if "anexo" in nombre_plantilla.lower() else 8
    
    # Procesar todos los párrafos del documento
    for p in doc.paragraphs:
        aplicar_estilo_y_reemplazo(p, datos, fuente_size)
    
    # Procesar tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p_celda in celda.paragraphs:
                    aplicar_estilo_y_reemplazo(p_celda, datos, fuente_size)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- INTERFAZ ---
st.title("📄 GENERADOR RÁPIDO")
st.subheader("Certificación de Ingresos")
st.info("📱 Ideal para uso desde el celular. Llena los datos abajo.")

# Verificar si existen los archivos antes de empezar
archivos_presentes = [p for p in PLANTILLAS_CONFIG if os.path.exists(p)]

if not archivos_presentes:
    st.error("❌ No se encontraron las plantillas .docx en el servidor.")
    st.write("Asegúrate de que 'certificacion.docx' y 'anexo.docx' estén en la misma carpeta.")
else:
    # Extraer etiquetas de los archivos locales automáticamente
    lista_unica = extraer_etiquetas_fijas(archivos_presentes)

    if lista_unica:
        # Formulario optimizado para móvil
        with st.form("datos_ingreso"):
            datos_usuario = {}
            for etiqueta in lista_unica:
                label = etiqueta.replace("_", " ").upper()
                datos_usuario[etiqueta] = st.text_input(label, placeholder=f"Escribe {label.lower()}...")
            
            st.write("---")
            boton_generar = st.form_submit_button("🚀 GENERAR DOCUMENTOS", use_container_width=True)

        if boton_generar:
            st.success("✨ Documentos listos para descargar:")
            
            for p_nombre in archivos_presentes:
                try:
                    archivo_final = generar_documento(p_nombre, datos_usuario)
                    
                    # Mostrar info del tamaño aplicado
                    size_info = "Arial 12" if "anexo" in p_nombre.lower() else "Arial 8"
                    
                    st.download_button(
                        label=f"⬇️ DESCARGAR {p_nombre.replace('.docx', '').upper()} ({size_info})",
                        data=archivo_final,
                        file_name=f"LISTO_{p_nombre}",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"Error al procesar {p_nombre}: {e}")
    else:
        st.warning("No se encontraron etiquetas {{...}} en las plantillas.")
